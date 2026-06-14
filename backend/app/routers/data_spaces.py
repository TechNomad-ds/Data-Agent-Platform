"""数据空间路由 - CRUD、文件关联、索引管理"""
import uuid
import io
import shutil
import zipfile
from pathlib import Path

import aiofiles

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File as FastAPIFile
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func

from app.core.database import get_db
from app.deps import get_current_user
from app.models.user import User
from app.models.file import File
from app.models.data_space import DataSpace, DataSpaceFile
from app.schemas.data_space import (
    DataSpaceCreate, DataSpaceUpdate, DataSpaceResponse,
    DataSpaceDetailResponse, FileInSpace, AddFilesRequest,
)
from app.schemas.file import FileResponse
from app.config import settings
from app.routers.files import ALLOWED_EXTENSIONS, MAX_FILE_SIZE, get_file_type, get_user_storage_path, _is_junk_path

router = APIRouter()


@router.post("", response_model=DataSpaceResponse, status_code=201)
async def create_data_space(
    data: DataSpaceCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """创建数据空间"""
    # 检查数量上限（管理员不限）
    if current_user.role != "admin":
        count_result = await db.execute(
            select(func.count()).select_from(DataSpace).where(DataSpace.user_id == current_user.id)
        )
        if (count_result.scalar() or 0) >= settings.max_spaces_per_user:
            raise HTTPException(status_code=400, detail=f"数据空间数量已达上限({settings.max_spaces_per_user}个)")

    # 检查同名
    result = await db.execute(
        select(DataSpace).where(DataSpace.user_id == current_user.id, DataSpace.name == data.name)
    )
    if result.scalar_one_or_none():
        raise HTTPException(status_code=400, detail="已存在同名数据空间")

    space = DataSpace(user_id=current_user.id, name=data.name, description=data.description)
    db.add(space)
    await db.flush()

    return DataSpaceResponse(
        id=space.id, name=space.name, description=space.description,
        file_count=0,
        created_at=space.created_at, updated_at=space.updated_at,
    )


@router.get("", response_model=list[DataSpaceResponse])
async def list_data_spaces(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取当前用户的数据空间列表"""
    from sqlalchemy import literal_column
    file_count_subq = (
        select(func.count())
        .select_from(DataSpaceFile)
        .where(DataSpaceFile.data_space_id == DataSpace.id)
        .correlate(DataSpace)
        .scalar_subquery()
    )
    result = await db.execute(
        select(DataSpace, file_count_subq.label("file_count"))
        .where(DataSpace.user_id == current_user.id)
        .order_by(DataSpace.updated_at.desc())
    )

    return [
        DataSpaceResponse(
            id=space.id, name=space.name, description=space.description,
            file_count=file_count or 0,
            created_at=space.created_at, updated_at=space.updated_at,
        )
        for space, file_count in result.all()
    ]


@router.get("/{space_id}", response_model=DataSpaceDetailResponse)
async def get_data_space(
    space_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取数据空间详情（含文件列表）"""
    result = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    space = result.scalar_one_or_none()
    if not space:
        raise HTTPException(status_code=404, detail="数据空间不存在")

    # 查询关联文件
    file_result = await db.execute(
        select(File, DataSpaceFile.added_at)
        .join(DataSpaceFile, DataSpaceFile.file_id == File.id)
        .where(DataSpaceFile.data_space_id == space_id)
        .order_by(DataSpaceFile.added_at.desc())
    )
    files_in_space = [
        FileInSpace(
            file_id=f.id, filename=f.filename,
            file_type=f.file_type, file_size=f.file_size, added_at=added_at,
        )
        for f, added_at in file_result.all()
    ]

    return DataSpaceDetailResponse(
        id=space.id, name=space.name, description=space.description,
        file_count=len(files_in_space),
        created_at=space.created_at, updated_at=space.updated_at,
        files=files_in_space,
    )


@router.put("/{space_id}", response_model=DataSpaceResponse)
async def update_data_space(
    space_id: uuid.UUID,
    data: DataSpaceUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """更新数据空间"""
    result = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    space = result.scalar_one_or_none()
    if not space:
        raise HTTPException(status_code=404, detail="数据空间不存在")

    if data.name is not None:
        space.name = data.name
    if data.description is not None:
        space.description = data.description

    count_result = await db.execute(
        select(func.count()).select_from(DataSpaceFile).where(DataSpaceFile.data_space_id == space.id)
    )
    file_count = count_result.scalar()

    return DataSpaceResponse(
        id=space.id, name=space.name, description=space.description,
        file_count=file_count,
        created_at=space.created_at, updated_at=space.updated_at,
    )


@router.delete("/{space_id}", status_code=204)
async def delete_data_space(
    space_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """删除数据空间"""
    result = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    space = result.scalar_one_or_none()
    if not space:
        raise HTTPException(status_code=404, detail="数据空间不存在")

    import logging
    _logger = logging.getLogger("data_spaces")

    # 清理 ChromaDB collection
    try:
        from app.services.embedding import get_chroma_client
        client = get_chroma_client()
        col_name = f"space_{str(space_id).replace('-', '')}"
        try:
            client.delete_collection(col_name)
        except Exception:
            pass
    except Exception as e:
        _logger.warning(f"清理 ChromaDB 失败: {e}")

    # 清理 SQLite 缓存
    try:
        from app.services.sqlite_engine import invalidate_cache
        invalidate_cache(str(space_id))
    except Exception as e:
        _logger.warning(f"清理 SQLite 缓存失败: {e}")

    # 清理检索缓存
    try:
        from app.services.retrieval import invalidate_cache as invalidate_retrieval_cache
        invalidate_retrieval_cache(str(space_id))
    except Exception as e:
        _logger.warning(f"清理检索缓存失败: {e}")

    await db.delete(space)


@router.post("/{space_id}/files", status_code=201)
async def add_files_to_space(
    space_id: uuid.UUID,
    data: AddFilesRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """向数据空间添加文件"""
    # 验证数据空间归属
    result = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="数据空间不存在")

    # 验证文件归属并添加
    for file_id in data.file_ids:
        file_result = await db.execute(
            select(File).where(File.id == file_id, File.user_id == current_user.id)
        )
        if not file_result.scalar_one_or_none():
            raise HTTPException(status_code=404, detail=f"文件 {file_id} 不存在")

        # 检查是否已关联
        existing = await db.execute(
            select(DataSpaceFile).where(
                DataSpaceFile.data_space_id == space_id, DataSpaceFile.file_id == file_id
            )
        )
        if not existing.scalar_one_or_none():
            db.add(DataSpaceFile(data_space_id=space_id, file_id=file_id))

    return {"message": "文件已添加到数据空间"}


@router.delete("/{space_id}/files/{file_id}", status_code=204)
async def remove_file_from_space(
    space_id: uuid.UUID,
    file_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """从数据空间移除文件"""
    result = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="数据空间不存在")

    link_result = await db.execute(
        select(DataSpaceFile).where(
            DataSpaceFile.data_space_id == space_id, DataSpaceFile.file_id == file_id
        )
    )
    link = link_result.scalar_one_or_none()
    if link:
        await db.delete(link)


@router.post("/{space_id}/upload", response_model=list[FileResponse], status_code=201)
async def upload_files_to_space(
    space_id: uuid.UUID,
    files: list[UploadFile] = FastAPIFile(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """上传文件到数据空间（上传并自动关联，ZIP 自动解压后每个文件单独入库）"""
    result = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    space = result.scalar_one_or_none()
    if not space:
        raise HTTPException(status_code=404, detail="数据空间不存在")

    # 检查文件数量上限（管理员不限）
    if current_user.role != "admin":
        file_count_result = await db.execute(
            select(func.count()).select_from(DataSpaceFile).where(DataSpaceFile.data_space_id == space_id)
        )
        current_file_count = file_count_result.scalar() or 0
        if current_file_count >= settings.max_files_per_space:
            raise HTTPException(status_code=400, detail=f"该数据空间文件数已达上限({settings.max_files_per_space}个)")

    if len(files) > 20:
        raise HTTPException(status_code=400, detail="单次最多上传 20 个文件")

    uploaded = []
    user_storage = get_user_storage_path(current_user.id)

    for upload_file in files:
        file_type = get_file_type(upload_file.filename)
        if file_type not in ALLOWED_EXTENSIONS:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file_type}")

        # 流式写入临时文件，避免大文件占满内存
        temp_path = user_storage / f"_upload_tmp_{uuid.uuid4()}"
        file_size = 0
        try:
            async with aiofiles.open(temp_path, "wb") as tmp:
                while chunk := await upload_file.read(1024 * 1024):
                    file_size += len(chunk)
                    if file_size > MAX_FILE_SIZE:
                        raise HTTPException(status_code=400, detail=f"文件 {upload_file.filename} 超过大小限制(200MB)")
                    await tmp.write(chunk)
        except HTTPException:
            temp_path.unlink(missing_ok=True)
            raise

        # ZIP 文件：解压后每个子文件单独入库并关联到数据空间
        if file_type == "zip":
            MAX_ZIP_FILES = 1000
            MAX_ZIP_TOTAL_SIZE = 500 * 1024 * 1024  # 500MB 解压上限
            tmp_dir = user_storage / f"_zip_tmp_{uuid.uuid4()}"
            tmp_dir.mkdir(parents=True, exist_ok=True)
            try:
                zf = zipfile.ZipFile(str(temp_path), "r")
                # 安全检查：路径遍历、文件数量、总大小
                # 计数只针对真正会入库的文件（排除目录、垃圾文件、不支持的扩展名），
                # 避免代码仓库类 zip 里大量无关文件把计数撑爆。
                total_uncompressed = 0
                valid_count = 0
                for info in zf.infolist():
                    if info.filename.startswith('/') or '..' in info.filename:
                        shutil.rmtree(tmp_dir, ignore_errors=True)
                        raise HTTPException(status_code=400, detail=f"zip 文件包含不安全的路径: {info.filename}")
                    if info.is_dir() or _is_junk_path(info.filename):
                        continue
                    if get_file_type(info.filename) not in ALLOWED_EXTENSIONS:
                        continue
                    valid_count += 1
                    total_uncompressed += info.file_size
                if valid_count > MAX_ZIP_FILES:
                    shutil.rmtree(tmp_dir, ignore_errors=True)
                    raise HTTPException(status_code=400, detail=f"zip 文件包含过多有效文件({valid_count}个，上限{MAX_ZIP_FILES}个)")
                if total_uncompressed > MAX_ZIP_TOTAL_SIZE:
                    shutil.rmtree(tmp_dir, ignore_errors=True)
                    raise HTTPException(status_code=400, detail=f"zip 解压后总大小超过限制(500MB)")
                zf.extractall(tmp_dir)
                member_names = zf.namelist()
                zf.close()
            except zipfile.BadZipFile:
                shutil.rmtree(tmp_dir, ignore_errors=True)
                raise HTTPException(status_code=400, detail=f"无效的 zip 文件: {upload_file.filename}")

            for member in member_names:
                if member.endswith("/") or _is_junk_path(member):
                    continue
                member_type = get_file_type(member)
                if member_type not in ALLOWED_EXTENSIONS:
                    continue

                src = tmp_dir / member
                if not src.is_file():
                    continue

                member_content = src.read_bytes()
                member_name = Path(member).name
                member_id = uuid.uuid4()
                member_dir = user_storage / str(member_id)
                member_dir.mkdir(parents=True, exist_ok=True)
                member_path = member_dir / member_name
                shutil.move(str(src), str(member_path))

                relative_path = str(member_path.relative_to(Path(settings.storage_root)))
                file_record = File(
                    id=member_id,
                    user_id=current_user.id,
                    filename=member_name,
                    original_filename=member_name,
                    file_type=member_type,
                    file_size=len(member_content),
                    storage_path=relative_path,
                    metadata_={"source_zip": upload_file.filename, "zip_path": member},
                )
                db.add(file_record)
                await db.flush()
                db.add(DataSpaceFile(data_space_id=space_id, file_id=member_id))
                uploaded.append(file_record)

            shutil.rmtree(tmp_dir, ignore_errors=True)
            temp_path.unlink(missing_ok=True)
            continue

        # 普通文件：直接 move 临时文件到目标位置（不读回内存）
        file_id = uuid.uuid4()
        file_dir = user_storage / str(file_id)
        file_dir.mkdir(parents=True, exist_ok=True)
        file_path = file_dir / upload_file.filename

        shutil.move(str(temp_path), str(file_path))

        relative_path = str(file_path.relative_to(Path(settings.storage_root)))
        file_record = File(
            id=file_id,
            user_id=current_user.id,
            filename=upload_file.filename,
            original_filename=upload_file.filename,
            file_type=file_type,
            file_size=file_size,
            storage_path=relative_path,
            mime_type=upload_file.content_type,
        )
        db.add(file_record)
        await db.flush()
        db.add(DataSpaceFile(data_space_id=space_id, file_id=file_id))
        uploaded.append(file_record)

    await db.flush()

    # 后台异步预处理
    import asyncio
    import logging
    from app.services.preprocessing import preprocess_file_limited, run_limited

    logger = logging.getLogger("data_spaces")

    async def _run_preprocessing():
        for f in uploaded:
            try:
                await preprocess_file_limited(f.id, space_id)
            except Exception as e:
                logger.error(f"预处理文件 {f.filename} 失败: {e}", exc_info=True)

        # 自动为文本/文档类文件构建知识图谱（无感化，不阻塞用户）
        if settings.graph_auto_extract:
            text_files = [f for f in uploaded if f.file_type in ("txt", "md", "pdf", "docx")]
            if text_files:
                try:
                    from app.services.graph import GraphService
                    gs = GraphService(str(current_user.id), str(space_id))
                    for f in text_files:
                        file_path = Path(settings.storage_root) / f.storage_path
                        if not file_path.exists():
                            continue
                        try:
                            if f.file_type in ("txt", "md"):
                                text = file_path.read_text(encoding="utf-8", errors="ignore")
                            elif f.file_type == "pdf":
                                import fitz
                                doc = fitz.open(str(file_path))
                                text = "\n".join(p.get_text() for p in doc)
                                doc.close()
                            elif f.file_type == "docx":
                                from docx import Document as DocxDoc
                                doc = DocxDoc(str(file_path))
                                text = "\n".join(p.text for p in doc.paragraphs)
                            else:
                                continue
                            if len(text.strip()) >= 100:
                                await run_limited(
                                    "graph",
                                    gs.extract_triples_from_text(
                                        text[:5000],
                                        max_triples=settings.graph_max_triples_per_file,
                                    ),
                                )
                        except Exception as e:
                            logger.warning(f"图谱抽取跳过 {f.filename}: {e}")
                except Exception as e:
                    logger.error(f"图谱自动构建异常: {e}", exc_info=True)

    asyncio.create_task(_run_preprocessing())

    return uploaded


@router.get("/{space_id}/processing-status")
async def get_processing_status(
    space_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取数据空间的文件处理状态"""
    from app.models.data_profile import DataProfile
    result = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="数据空间不存在")

    # 统计文件总数
    file_count_result = await db.execute(
        select(func.count()).select_from(DataSpaceFile).where(DataSpaceFile.data_space_id == space_id)
    )
    total_files = file_count_result.scalar()

    # 统计已完成的 profile 数
    profile_result = await db.execute(
        select(DataProfile).where(DataProfile.data_space_id == space_id)
    )
    profiles = profile_result.scalars().all()
    ready = sum(1 for p in profiles if p.status == "ready")
    processing = sum(1 for p in profiles if p.status in ("pending", "processing"))
    error = sum(1 for p in profiles if p.status == "error")
    no_profile = total_files - len(profiles)

    return {
        "total_files": total_files,
        "ready": ready,
        "processing": processing + no_profile,
        "error": error,
        "all_ready": ready == total_files and total_files > 0,
    }


@router.get("/{space_id}/profile")
async def get_space_profile_endpoint(
    space_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取数据空间的聚合数据画像"""
    result = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="数据空间不存在")

    from app.services.preprocessing import get_space_profile
    profile = await get_space_profile(space_id)
    return profile


@router.get("/{space_id}/files/{file_id}/profile")
async def get_file_profile_endpoint(
    space_id: uuid.UUID,
    file_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取单个文件的数据画像"""
    # 先校验数据空间归属，防止越权读取他人文件画像（IDOR）
    space_check = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    if not space_check.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="数据空间不存在")

    from app.models.data_profile import DataProfile as DP
    result = await db.execute(
        select(DP).where(DP.file_id == file_id, DP.data_space_id == space_id)
    )
    profile = result.scalar_one_or_none()
    if not profile:
        raise HTTPException(status_code=404, detail="画像不存在，可能正在处理中")
    return {
        "file_id": str(profile.file_id),
        "profile_type": profile.profile_type,
        "status": profile.status,
        "profile_data": profile.profile_data,
        "error_message": profile.error_message,
    }


@router.get("/{space_id}/files/{file_id}/preview")
async def preview_file_data(
    space_id: uuid.UUID,
    file_id: uuid.UUID,
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=50, ge=10, le=200),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """预览文件数据（表格形式，支持分页）"""
    import pandas as pd
    import json as json_mod

    result = await db.execute(
        select(DataSpace).where(DataSpace.id == space_id, DataSpace.user_id == current_user.id)
    )
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="数据空间不存在")

    file_result = await db.execute(
        select(File).where(File.id == file_id, File.user_id == current_user.id)
    )
    file = file_result.scalar_one_or_none()
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")

    file_path = Path(settings.storage_root) / file.storage_path
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不在磁盘上")

    ext = file.file_type.lower()

    # SQLite database files
    if ext in ("sqlite", "db", "sqlite3"):
        import sqlite3
        conn = sqlite3.connect(str(file_path))
        try:
            cursor = conn.execute("SELECT name FROM sqlite_master WHERE type='table'")
            tables = [row[0] for row in cursor.fetchall()]
            if not tables:
                return {"type": "database", "tables": [], "message": "数据库为空"}

            target_table = tables[0]
            # 使用参数化的方式安全引用表名（SQLite 不支持参数化表名，用白名单校验）
            if target_table not in tables:
                return {"type": "database", "tables": tables, "message": "表不存在"}

            quoted_table = '"' + target_table.replace('"', '""') + '"'
            count_cursor = conn.execute(f"SELECT COUNT(*) FROM {quoted_table}")
            total = count_cursor.fetchone()[0]

            offset = (page - 1) * page_size
            data_cursor = conn.execute(f"SELECT * FROM {quoted_table} LIMIT ? OFFSET ?", (page_size, offset))
            col_names = [desc[0] for desc in data_cursor.description]
            rows_data = [list(str(v) if v is not None else "" for v in row) for row in data_cursor.fetchall()]

            return {
                "type": "database",
                "tables": tables,
                "current_table": target_table,
                "columns": [{"name": c, "dtype": "text"} for c in col_names],
                "rows": rows_data,
                "total_rows": total,
                "page": page,
                "page_size": page_size,
                "filename": file.filename,
            }
        finally:
            conn.close()

    # Tabular formats
    df = None
    if ext == "csv":
        from app.services.preprocessing import _detect_encoding
        encoding = _detect_encoding(file_path)
        df = pd.read_csv(file_path, encoding=encoding, nrows=page * page_size, on_bad_lines="skip")
    elif ext == "tsv":
        from app.services.preprocessing import _detect_encoding
        encoding = _detect_encoding(file_path)
        df = pd.read_csv(file_path, sep="\t", encoding=encoding, nrows=page * page_size, on_bad_lines="skip")
    elif ext in ("xlsx", "xls"):
        df = pd.read_excel(file_path, nrows=page * page_size)
    elif ext == "parquet":
        df = pd.read_parquet(file_path).head(page * page_size)
    elif ext == "feather":
        df = pd.read_feather(file_path).head(page * page_size)
    elif ext == "json":
        content = file_path.read_text(encoding="utf-8")
        data = json_mod.loads(content)
        if isinstance(data, list):
            df = pd.DataFrame(data[:page * page_size])
        elif isinstance(data, dict) and "records" in data:
            df = pd.DataFrame(data["records"][:page * page_size])
        else:
            return {"type": "text", "content": content[:2000], "total_chars": len(content)}
    elif ext == "jsonl":
        df = pd.read_json(file_path, lines=True, nrows=page * page_size)
    elif ext == "dta":
        df = pd.read_stata(file_path).head(page * page_size)
    elif ext == "sav":
        df = pd.read_spss(file_path).head(page * page_size)
    elif ext == "sas7bdat":
        df = pd.read_sas(file_path).head(page * page_size)
    elif ext in ("txt", "md", "py", "sql", "html", "xml", "yaml", "yml", "log", "r", "ipynb"):
        from app.services.preprocessing import _detect_encoding
        encoding = _detect_encoding(file_path)
        content = file_path.read_text(encoding=encoding, errors="ignore")
        lines = content.split("\n")
        start = (page - 1) * page_size
        end = start + page_size
        return {
            "type": "text",
            "content": "\n".join(lines[start:end]),
            "total_lines": len(lines),
            "page": page,
            "page_size": page_size,
        }
    elif ext == "pdf":
        try:
            import fitz
            doc = fitz.open(str(file_path))
            text_parts = []
            for i, pg in enumerate(doc):
                text_parts.append(pg.get_text())
            doc.close()
            full_text = "\n\n".join(text_parts)
            lines = full_text.split("\n")
            start = (page - 1) * page_size
            end = start + page_size
            return {
                "type": "text",
                "content": "\n".join(lines[start:end]),
                "total_lines": len(lines),
                "page": page,
                "page_size": page_size,
                "filename": file.filename,
            }
        except ImportError:
            return {"type": "unsupported", "message": "需要安装 PyMuPDF 才能预览 PDF"}
    elif ext == "docx":
        try:
            from docx import Document
            doc = Document(str(file_path))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            lines = full_text.split("\n")
            start = (page - 1) * page_size
            end = start + page_size
            return {
                "type": "text",
                "content": "\n".join(lines[start:end]),
                "total_lines": len(lines),
                "page": page,
                "page_size": page_size,
                "filename": file.filename,
            }
        except ImportError:
            return {"type": "unsupported", "message": "需要安装 python-docx 才能预览 Word"}
    elif ext in ("png", "jpg", "jpeg", "gif", "bmp", "webp"):
        return {"type": "image", "filename": file.filename, "file_size_kb": round(file_path.stat().st_size / 1024, 1)}
    else:
        return {"type": "unsupported", "message": f"不支持预览的文件类型: {ext}"}

    if df is None or df.empty:
        return {"type": "unsupported", "message": "无法解析文件内容"}

    total_rows = len(df)
    start = (page - 1) * page_size
    end = start + page_size
    page_df = df.iloc[start:end]

    columns = []
    for col in df.columns:
        columns.append({
            "name": str(col),
            "dtype": str(df[col].dtype),
        })

    rows = page_df.fillna("").astype(str).values.tolist()

    return {
        "type": "table",
        "columns": columns,
        "rows": rows,
        "total_rows": total_rows,
        "page": page,
        "page_size": page_size,
        "filename": file.filename,
    }
