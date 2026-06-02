"""数据预处理服务 - 文件上传后自动分析生成数据画像"""
import uuid
import json
from pathlib import Path
from typing import Any, Dict, Optional

import pandas as pd
import numpy as np
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.core.database import get_session_factory
from app.models.file import File
from app.models.data_profile import DataProfile
from app.services.chunking import greedy_chunk
from app.services import embedding as embed_svc


async def preprocess_file(file_id: uuid.UUID, data_space_id: uuid.UUID) -> Dict[str, Any]:
    """预处理单个文件，生成数据画像。向量索引在后台异步完成，不阻塞状态。"""
    async with get_session_factory()() as db:
        result = await db.execute(select(File).where(File.id == file_id))
        file = result.scalar_one_or_none()
        if not file:
            return {"error": "文件不存在"}

        file_path = Path(settings.storage_root) / file.storage_path
        if not file_path.exists():
            return {"error": "文件不在磁盘上"}

        ext = file.file_type.lower()

        try:
            if ext in ("csv", "tsv", "xlsx", "xls", "json", "jsonl", "parquet", "feather", "dta", "sav", "sas7bdat"):
                profile_data = _profile_tabular(file_path, ext)
                profile_type = "tabular"
                import asyncio
                asyncio.create_task(
                    _embed_tabular_background(profile_data, str(file_id), str(data_space_id), file.filename)
                )
            elif ext in ("sqlite", "db", "sqlite3"):
                profile_data = _profile_sqlite(file_path)
                profile_type = "database"
            elif ext in ("txt", "md", "py", "sql", "html", "xml", "yaml", "yml", "log", "r", "ipynb"):
                profile_data = _profile_text_fast(file_path)
                profile_type = "text"
                import asyncio
                asyncio.create_task(_embed_text_background(file_path, str(file_id), str(data_space_id), file.filename))
            elif ext in ("pdf", "docx"):
                profile_data = _profile_document_fast(file_path, ext)
                profile_type = "document"
                import asyncio
                asyncio.create_task(_embed_document_background(file_path, ext, str(file_id), str(data_space_id), file.filename))
            elif ext in ("png", "jpg", "jpeg", "gif", "bmp", "webp"):
                profile_data = _profile_image(file_path)
                profile_type = "image"
            else:
                profile_data = {"info": f"文件类型: {ext}", "file_size": file.file_size}
                profile_type = "other"

            existing = await db.execute(
                select(DataProfile).where(
                    DataProfile.file_id == file_id,
                    DataProfile.data_space_id == data_space_id,
                )
            )
            profile = existing.scalar_one_or_none()
            if profile:
                profile.profile_data = profile_data
                profile.profile_type = profile_type
                profile.status = "ready"
                profile.error_message = None
            else:
                profile = DataProfile(
                    file_id=file_id,
                    data_space_id=data_space_id,
                    profile_type=profile_type,
                    profile_data=profile_data,
                    status="ready",
                )
                db.add(profile)

            await db.commit()
            return profile_data

        except Exception as e:
            existing = await db.execute(
                select(DataProfile).where(
                    DataProfile.file_id == file_id,
                    DataProfile.data_space_id == data_space_id,
                )
            )
            profile = existing.scalar_one_or_none()
            if profile:
                profile.status = "error"
                profile.error_message = str(e)
            else:
                db.add(DataProfile(
                    file_id=file_id,
                    data_space_id=data_space_id,
                    profile_type="tabular",
                    profile_data={},
                    status="error",
                    error_message=str(e),
                ))
            await db.commit()
            return {"error": str(e)}


def _profile_tabular(file_path: Path, ext: str) -> Dict[str, Any]:
    """生成表格数据画像 - 支持多种格式"""
    df = _load_tabular(file_path, ext)
    if df is None or df.empty:
        return {"error": f"无法加载 {ext} 格式文件"}

    profile: Dict[str, Any] = {
        "row_count": int(len(df)),
        "column_count": int(len(df.columns)),
        "columns": [],
        "memory_usage_mb": round(df.memory_usage(deep=True).sum() / 1024 / 1024, 2),
    }

    numeric_cols = []
    for col in df.columns:
        col_info: Dict[str, Any] = {
            "name": col,
            "dtype": str(df[col].dtype),
            "non_null_count": int(df[col].notna().sum()),
            "null_count": int(df[col].isna().sum()),
            "null_pct": round(float(df[col].isna().mean()) * 100, 1),
            "unique_count": int(df[col].nunique()),
        }

        if pd.api.types.is_numeric_dtype(df[col]):
            desc = df[col].describe()
            col_info["stats"] = {
                "mean": _safe_float(desc.get("mean")),
                "std": _safe_float(desc.get("std")),
                "min": _safe_float(desc.get("min")),
                "25%": _safe_float(desc.get("25%")),
                "50%": _safe_float(desc.get("50%")),
                "75%": _safe_float(desc.get("75%")),
                "max": _safe_float(desc.get("max")),
            }
            numeric_cols.append(col)
        else:
            top_values = df[col].value_counts().head(10)
            col_info["top_values"] = {str(k): int(v) for k, v in top_values.items()}

        col_info["sample_values"] = [str(v) for v in df[col].dropna().head(5).tolist()]
        profile["columns"].append(col_info)

    # Data quality metrics
    quality = {
        "duplicate_rows": int(df.duplicated().sum()),
        "duplicate_pct": round(float(df.duplicated().mean()) * 100, 1),
        "complete_rows": int((~df.isna().any(axis=1)).sum()),
        "complete_pct": round(float((~df.isna().any(axis=1)).mean()) * 100, 1),
    }

    # Outlier detection for numeric columns (IQR method)
    outlier_cols = []
    for col in numeric_cols[:10]:
        q1 = df[col].quantile(0.25)
        q3 = df[col].quantile(0.75)
        iqr = q3 - q1
        if iqr > 0:
            outlier_count = int(((df[col] < q1 - 1.5 * iqr) | (df[col] > q3 + 1.5 * iqr)).sum())
            if outlier_count > 0:
                outlier_cols.append({"column": col, "outlier_count": outlier_count, "outlier_pct": round(outlier_count / len(df) * 100, 1)})

    quality["outlier_columns"] = outlier_cols

    # Type suggestions
    type_suggestions = []
    for col in df.columns:
        if df[col].dtype == "object":
            non_null = df[col].dropna()
            if len(non_null) > 0:
                try:
                    pd.to_numeric(non_null.head(100))
                    type_suggestions.append({"column": col, "suggestion": "可转为数值类型"})
                    continue
                except (ValueError, TypeError):
                    pass
                try:
                    pd.to_datetime(non_null.head(100))
                    type_suggestions.append({"column": col, "suggestion": "可转为日期类型"})
                except (ValueError, TypeError):
                    pass

    quality["type_suggestions"] = type_suggestions
    profile["quality"] = quality

    if len(numeric_cols) >= 2 and len(numeric_cols) <= 20:
        corr = df[numeric_cols].corr()
        profile["correlations"] = {
            "columns": numeric_cols,
            "matrix": corr.fillna(0).values.tolist(),
        }

    return profile


def _profile_text_fast(file_path: Path) -> Dict[str, Any]:
    """快速生成文本文件画像（不做 embedding）"""
    content = file_path.read_text(encoding="utf-8", errors="ignore")
    lines = content.split("\n")
    return {
        "char_count": len(content),
        "line_count": len(lines),
        "word_count": len(content.split()),
        "preview": content[:500],
    }


async def _embed_tabular_background(
    profile_data: Dict[str, Any], file_id: str, data_space_id: str, filename: str
) -> None:
    """将表格文件的列信息和摘要嵌入为可搜索的文本块"""
    import logging
    logger = logging.getLogger("preprocessing")
    try:
        chunks = []
        columns = profile_data.get("columns", [])
        row_count = profile_data.get("row_count", 0)

        col_names = [c["name"] for c in columns]
        summary = f"表格文件 {filename}，共 {row_count} 行，{len(columns)} 列。列包括: {', '.join(col_names)}。"

        for c in columns:
            desc = f"列 {c['name']} (类型: {c['dtype']})"
            if c.get("stats"):
                s = c["stats"]
                desc += f"，均值 {s.get('mean', '?')}，范围 {s.get('min', '?')} ~ {s.get('max', '?')}"
            if c.get("top_values"):
                top = list(c["top_values"].keys())[:5]
                desc += f"，常见值: {', '.join(top)}"
            if c.get("sample_values"):
                desc += f"，示例: {', '.join(c['sample_values'][:3])}"
            summary += " " + desc + "。"

        chunks.append({"text": summary, "start_char": 0, "end_char": len(summary)})
        embed_svc.embed_chunks(data_space_id, chunks, file_id, filename)

        try:
            from app.services.retrieval import invalidate_cache
            invalidate_cache(data_space_id)
        except Exception:
            pass
        logger.info(f"表格嵌入完成: {filename}")
    except Exception as e:
        logger.error(f"表格嵌入失败 ({filename}): {e}", exc_info=True)


async def _embed_text_background(
    file_path: Path, file_id: str, data_space_id: str, filename: str
) -> None:
    """后台异步执行文本 embedding"""
    import logging
    logger = logging.getLogger("preprocessing")
    try:
        content = file_path.read_text(encoding="utf-8", errors="ignore")
        chunks = greedy_chunk(content, max_size=1000, overlap=200)
        embed_svc.embed_chunks(data_space_id, chunks, file_id, filename)
        try:
            from app.services.retrieval import invalidate_cache
            invalidate_cache(data_space_id)
        except Exception:
            pass
        logger.info(f"文本嵌入完成: {filename}, {len(chunks)} 个块")
    except Exception as e:
        logger.error(f"文本嵌入失败 ({filename}): {e}", exc_info=True)


def _profile_document_fast(file_path: Path, ext: str) -> Dict[str, Any]:
    """快速生成文档画像（不做 embedding）"""
    text = ""
    page_count = 0

    if ext == "pdf":
        try:
            import fitz
            doc = fitz.open(str(file_path))
            page_count = len(doc)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
        except ImportError:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
    elif ext == "docx":
        try:
            from docx import Document
            doc = Document(str(file_path))
            for para in doc.paragraphs:
                text += para.text + "\n"
            page_count = len(doc.paragraphs) // 30 + 1
        except ImportError:
            text = file_path.read_text(encoding="utf-8", errors="ignore")

    return {
        "char_count": len(text),
        "page_count": page_count,
        "preview": text[:500],
    }


async def _embed_document_background(
    file_path: Path, ext: str, file_id: str, data_space_id: str, filename: str
) -> None:
    """后台异步执行文档 embedding"""
    import logging
    logger = logging.getLogger("preprocessing")
    try:
        text = ""
        if ext == "pdf":
            try:
                import fitz
                doc = fitz.open(str(file_path))
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except ImportError:
                text = file_path.read_text(encoding="utf-8", errors="ignore")
        elif ext == "docx":
            try:
                from docx import Document
                doc = Document(str(file_path))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except ImportError:
                text = file_path.read_text(encoding="utf-8", errors="ignore")

        chunks = greedy_chunk(text, max_size=1000, overlap=200)
        embed_svc.embed_chunks(data_space_id, chunks, file_id, filename)
        try:
            from app.services.retrieval import invalidate_cache
            invalidate_cache(data_space_id)
        except Exception:
            pass
        logger.info(f"文档嵌入完成: {filename}, {len(chunks)} 个块")
    except Exception as e:
        logger.error(f"文档嵌入失败 ({filename}): {e}", exc_info=True)


def _load_json_df(file_path: Path) -> pd.DataFrame:
    content = file_path.read_text(encoding="utf-8")
    data = json.loads(content)
    if isinstance(data, list):
        return pd.DataFrame(data)
    elif isinstance(data, dict) and "records" in data:
        return pd.DataFrame(data["records"])
    elif isinstance(data, dict):
        return pd.DataFrame([data])
    return pd.DataFrame()


def _detect_encoding(file_path: Path) -> str:
    """检测文件编码"""
    try:
        import chardet
        raw = file_path.read_bytes()[:10000]
        result = chardet.detect(raw)
        return result.get("encoding", "utf-8") or "utf-8"
    except ImportError:
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1", "shift_jis"):
            try:
                file_path.read_text(encoding=enc)
                return enc
            except (UnicodeDecodeError, LookupError):
                continue
        return "utf-8"


def _load_tabular(file_path: Path, ext: str) -> Optional[pd.DataFrame]:
    """加载各种格式的表格数据"""
    try:
        if ext == "csv":
            encoding = _detect_encoding(file_path)
            return pd.read_csv(file_path, encoding=encoding, nrows=10000, on_bad_lines="skip")
        elif ext == "tsv":
            encoding = _detect_encoding(file_path)
            return pd.read_csv(file_path, sep="\t", encoding=encoding, nrows=10000, on_bad_lines="skip")
        elif ext in ("xlsx", "xls"):
            return pd.read_excel(file_path, nrows=10000)
        elif ext == "json":
            return _load_json_df(file_path)
        elif ext == "jsonl":
            return pd.read_json(file_path, lines=True, nrows=10000)
        elif ext == "parquet":
            return pd.read_parquet(file_path).head(10000)
        elif ext == "feather":
            return pd.read_feather(file_path).head(10000)
        elif ext == "dta":
            return pd.read_stata(file_path)
        elif ext == "sav":
            return pd.read_spss(file_path)
        elif ext == "sas7bdat":
            return pd.read_sas(file_path)
    except ImportError as e:
        import logging
        logging.getLogger("preprocessing").warning(f"缺少依赖库，无法加载 {ext}: {e}")
        return None
    except Exception:
        return None
    return None


def _profile_sqlite(file_path: Path) -> Dict[str, Any]:
    """分析 SQLite 数据库文件"""
    import sqlite3
    conn = sqlite3.connect(str(file_path))
    try:
        cursor = conn.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = []
        total_rows = 0
        for (table_name,) in cursor.fetchall():
            quoted = '"' + table_name.replace('"', '""') + '"'
            info_cursor = conn.execute(f"PRAGMA table_info({quoted})")
            columns = [{"name": row[1], "type": row[2], "nullable": not row[3]} for row in info_cursor.fetchall()]
            count_cursor = conn.execute(f"SELECT COUNT(*) FROM {quoted}")
            row_count = count_cursor.fetchone()[0]
            total_rows += row_count

            sample_cursor = conn.execute(f"SELECT * FROM {quoted} LIMIT 5")
            sample_rows = [list(row) for row in sample_cursor.fetchall()]

            tables.append({
                "name": table_name,
                "columns": columns,
                "row_count": row_count,
                "sample_rows": sample_rows,
            })

        return {
            "database_type": "sqlite",
            "table_count": len(tables),
            "total_rows": total_rows,
            "tables": tables,
            "file_size_mb": round(file_path.stat().st_size / 1024 / 1024, 2),
        }
    finally:
        conn.close()


def _profile_image(file_path: Path) -> Dict[str, Any]:
    """分析图片文件"""
    try:
        from PIL import Image
        img = Image.open(file_path)
        return {
            "width": img.width,
            "height": img.height,
            "format": img.format,
            "mode": img.mode,
            "file_size_kb": round(file_path.stat().st_size / 1024, 1),
        }
    except ImportError:
        return {
            "file_size_kb": round(file_path.stat().st_size / 1024, 1),
            "note": "需要安装 Pillow 库以获取图片详细信息",
        }
    except Exception as e:
        return {"error": str(e)}


def _safe_float(val) -> Optional[float]:
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return None
    return round(float(val), 4)


async def get_space_profile(data_space_id: uuid.UUID) -> Dict[str, Any]:
    """获取数据空间的聚合画像"""
    async with get_session_factory()() as db:
        result = await db.execute(
            select(DataProfile).where(
                DataProfile.data_space_id == data_space_id,
                DataProfile.status == "ready",
            )
        )
        profiles = result.scalars().all()

    if not profiles:
        return {"status": "empty", "files": []}

    files_summary = []
    total_rows = 0
    total_columns = 0

    for p in profiles:
        summary = {
            "file_id": str(p.file_id),
            "profile_type": p.profile_type,
        }
        if p.profile_type == "tabular":
            summary["row_count"] = p.profile_data.get("row_count", 0)
            summary["column_count"] = p.profile_data.get("column_count", 0)
            summary["columns"] = [c["name"] for c in p.profile_data.get("columns", [])]
            total_rows += summary["row_count"]
            total_columns += summary["column_count"]
        elif p.profile_type in ("text", "document"):
            summary["char_count"] = p.profile_data.get("char_count", 0)
            summary["chunk_count"] = p.profile_data.get("chunk_count", 0)

        files_summary.append(summary)

    return {
        "status": "ready",
        "file_count": len(profiles),
        "total_rows": total_rows,
        "total_columns": total_columns,
        "files": files_summary,
    }
